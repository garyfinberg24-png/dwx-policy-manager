"""
Generate PolicyIQ Business Proposal for Clicks as a professionally formatted Word document.
Run: python docs/generate-proposal-docx.py
Output: docs/PolicyIQ-Proposal-Clicks.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colours ──
TEAL = RGBColor(0x0D, 0x94, 0x88)
TEAL_DARK = RGBColor(0x0F, 0x76, 0x6E)
NAVY = RGBColor(0x0F, 0x17, 0x2A)
SLATE = RGBColor(0x33, 0x41, 0x55)
MUTED = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEAL_BG = "D1FAE5"  # light green-teal for table headers
PALE_BG = "F0FDFA"
LIGHT_GREY = "F8FAFC"
BORDER_CLR = "E2E8F0"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(10.5)
font.color.rgb = SLATE
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25


# ── Helper functions ──

def add_blank(count=1):
    for _ in range(count):
        doc.add_paragraph()

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "E2E8F0")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_teal_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else TEAL_DARK
        run.font.name = 'Segoe UI'
    return h

def add_section_header(number, label, title, subtitle=None):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}  //  {label.upper()}")
    run.font.size = Pt(9)
    run.font.color.rgb = TEAL
    run.font.bold = True
    run.font.name = 'Segoe UI'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(24)

    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = 'Segoe UI Semibold'
        run.font.size = Pt(22)
    h.paragraph_format.space_after = Pt(4)

    if subtitle:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(11)
        run2.font.color.rgb = MUTED
        run2.font.italic = True
        run2.font.name = 'Segoe UI'
        p2.paragraph_format.space_after = Pt(16)

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    # Handle bold markers **text**
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10.5)
        run.font.color.rgb = SLATE
        if i % 2 == 1:
            run.bold = True
            run.font.color.rgb = NAVY
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        run2 = p.add_run(f" {text}")
        run2.font.name = 'Segoe UI'
        run2.font.size = Pt(10)
        run2.font.color.rgb = SLATE
    else:
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10)
        run.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(3)
    return p

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"\u201C{text}\u201D")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL_DARK
    run.italic = True
    return p

def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header.upper())
        run.font.name = 'Segoe UI'
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = TEAL_DARK
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        set_cell_shading(cell, TEAL_BG)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9.5)
            run.font.color.rgb = SLATE
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if r % 2 == 1:
                set_cell_shading(cell, LIGHT_GREY)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table

def add_kpi_table(kpis):
    """Add a row of KPI cards as a table."""
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (value, label) in enumerate(kpis):
        # Value cell
        cell_v = table.rows[0].cells[i]
        cell_v.text = ''
        p = cell_v.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(value))
        run.font.name = 'Segoe UI'
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        set_cell_shading(cell_v, PALE_BG)

        # Label cell
        cell_l = table.rows[1].cells[i]
        cell_l.text = ''
        p2 = cell_l.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label.upper())
        run2.font.name = 'Segoe UI'
        run2.font.size = Pt(7.5)
        run2.font.bold = True
        run2.font.color.rgb = MUTED
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(12)
        set_cell_shading(cell_l, PALE_BG)

    doc.add_paragraph()
    return table

def add_sub_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Segoe UI Semibold'
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    run.bold = True
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_page_break():
    doc.add_page_break()

def add_signing_field(label):
    p = doc.add_paragraph()
    run = p.add_run(label.upper())
    run.font.name = 'Segoe UI'
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.bold = True
    p.paragraph_format.space_after = Pt(0)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(16)
    # Add a line using bottom border
    pPr = p2._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CBD5E1"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run2 = p2.add_run("  ")
    run2.font.size = Pt(14)


# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════

add_blank(5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DWx  DIGITAL WORKPLACE EXCELLENCE")
run.font.name = 'Segoe UI'
run.font.size = Pt(9)
run.font.color.rgb = TEAL
run.bold = True
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PolicyIQ")
run.font.name = 'Segoe UI'
run.font.size = Pt(48)
run.font.color.rgb = TEAL
run.bold = True
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("POLICY GOVERNANCE & COMPLIANCE")
run.font.name = 'Segoe UI'
run.font.size = Pt(11)
run.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(36)

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u2500" * 12)
run.font.color.rgb = TEAL
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(36)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Business Proposal for")
run.font.name = 'Segoe UI'
run.font.size = Pt(16)
run.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Clicks")
run.font.name = 'Segoe UI'
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.bold = True
p.paragraph_format.space_after = Pt(48)

# Meta info
meta_table = doc.add_table(rows=1, cols=3)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate([
    ("PREPARED BY", "First Digital"),
    ("DATE", "April 2026"),
    ("VERSION", "1.0 \u2014 Confidential")
]):
    cell = meta_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_l = p.add_run(f"{label}\n")
    run_l.font.name = 'Segoe UI'
    run_l.font.size = Pt(7.5)
    run_l.font.color.rgb = MUTED
    run_l.bold = True

    run_v = p.add_run(value)
    run_v.font.name = 'Segoe UI'
    run_v.font.size = Pt(10)
    run_v.font.color.rgb = NAVY

add_page_break()


# ═══════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════

add_section_header("01", "Executive Summary",
    "Transforming Policy Management for Clicks",
    "From manual processes to intelligent, automated governance")

add_body(
    "Clicks operates in a highly regulated retail and healthcare environment where policy compliance "
    "is not optional \u2014 it is a business imperative. From pharmaceutical dispensing protocols to "
    "employment practices, **every policy must be current, acknowledged, and auditable**."
)

add_body(
    "**PolicyIQ** is an enterprise-grade policy governance platform built on Microsoft 365 and "
    "SharePoint Online. It delivers end-to-end policy lifecycle management \u2014 from authoring and "
    "approval through distribution, acknowledgement, quiz-based comprehension testing, and compliance "
    "analytics \u2014 all within the Microsoft ecosystem Clicks already uses."
)

add_body(
    "With **AI-powered capabilities** including an intelligent policy assistant, automated quiz "
    "generation, and bulk document classification, PolicyIQ eliminates manual overhead while "
    "strengthening compliance posture across the entire organisation."
)

add_kpi_table([
    ("17", "Purpose-Built Views"),
    ("150+", "Business Services"),
    ("30+", "SharePoint Lists"),
    ("3", "AI-Powered Features"),
])


# ═══════════════════════════════════════════════════════════
# SECTION 2 — THE CHALLENGE
# ═══════════════════════════════════════════════════════════

add_section_header("02", "The Challenge",
    "Why Policy Management Matters for Clicks",
    "Retail, pharmacy, and healthcare compliance across a national footprint")

add_body(
    "Managing policies manually across hundreds of locations creates significant risk and overhead:"
)

add_bullet("Pharmaceutical, health & safety, and employment regulations demand provable compliance. "
           "Manual tracking creates audit gaps that carry legal and financial risk.",
           bold_prefix="Regulatory Exposure \u2014")

add_bullet("Email-based distribution, spreadsheet tracking, and paper-based acknowledgements "
           "consume countless hours from HR, compliance, and line managers.",
           bold_prefix="Manual Overhead \u2014")

add_bullet("Hundreds of store locations, pharmacies, and support offices mean policy changes "
           "must reach every employee \u2014 reliably, verifiably, on time.",
           bold_prefix="Distributed Workforce \u2014")

add_bullet("Multiple versions of policies across shared drives, intranets, and email attachments "
           "create confusion about which version is current and authoritative.",
           bold_prefix="Version Control Chaos \u2014")

add_quote(
    "The cost of non-compliance is always higher than the cost of compliance. PolicyIQ ensures "
    "Clicks has a single source of truth for every policy, every acknowledgement, and every audit "
    "trail \u2014 in real time."
)


# ═══════════════════════════════════════════════════════════
# SECTION 3 — VALUE PROPOSITION
# ═══════════════════════════════════════════════════════════

add_page_break()
add_section_header("03", "Value Proposition",
    "What PolicyIQ Delivers",
    "Enterprise governance made simple, intelligent, and auditable")

add_bullet("Every action is logged with immutable timestamps \u2014 who read what, when they "
           "acknowledged, quiz scores, exemptions, and every approval step.",
           bold_prefix="Complete Audit Trail \u2014")

add_bullet("Auto-distribution, escalating reminders, SLA tracking, and deadline management "
           "\u2014 no employee falls through the cracks.",
           bold_prefix="Automated Compliance \u2014")

add_bullet("Employees, authors, managers, and administrators each get a tailored interface "
           "designed for their specific responsibilities.",
           bold_prefix="Purpose-Built Dashboards \u2014")

add_bullet("Conversational policy assistant, automated quiz generation from policy content, "
           "and intelligent bulk document classification.",
           bold_prefix="AI-Powered Intelligence \u2014")

add_bullet("Built on SharePoint Online and Azure \u2014 no third-party platforms, no additional "
           "licensing. Leverages your existing M365 investment.",
           bold_prefix="Microsoft 365 Native \u2014")

add_bullet("Four-tier access model (Employee, Author, Manager, Admin) ensures every user sees "
           "exactly what they need \u2014 nothing more, nothing less.",
           bold_prefix="Role-Based Access \u2014")


# ═══════════════════════════════════════════════════════════
# SECTION 4 — KEY FEATURES
# ═══════════════════════════════════════════════════════════

add_page_break()
add_section_header("04", "Key Features",
    "Feature Overview",
    "A comprehensive platform built for enterprise policy governance")

add_sub_heading("Policy Lifecycle Management")

add_body("Every policy follows a structured, auditable lifecycle: "
         "**Draft** \u2192 **In Review** \u2192 **Approval** \u2192 **Published** \u2192 **Archived / Retired**")

add_bullet("Guided 8-step creation wizard covering metadata, audience targeting, review dates, "
           "approval workflow, and rich content editing with WYSIWYG editor",
           bold_prefix="Policy Wizard \u2014")

add_bullet("Author in HTML, upload Word / Excel / PowerPoint / PDF documents, or create directly "
           "in Office Online. Automatic HTML conversion at publish time.",
           bold_prefix="Multi-Format Support \u2014")

add_bullet("Major/minor versioning with side-by-side diff comparison, rollback capability, "
           "and complete version history",
           bold_prefix="Version Control \u2014")

add_bullet("Multi-level approval chains with configurable routing, escalation rules, "
           "and delegation support",
           bold_prefix="Approval Workflows \u2014")

add_bullet("Bundle related policies for onboarding, compliance training, or role-specific "
           "assignment with pack-level approval",
           bold_prefix="Policy Packs \u2014")

add_bullet("Drag-and-drop import of up to 50 documents with automatic metadata extraction, "
           "AI classification, and batch assignment",
           bold_prefix="Bulk Upload \u2014")

add_sub_heading("Acknowledgement & Compliance Tracking")

add_bullet("One-time, annual, quarterly, monthly, on-update, and conditional re-certification",
           bold_prefix="Flexible Acknowledgement Types \u2014")

add_bullet("11 question types including multiple choice, matching, ordering, fill-in-the-blank, "
           "hotspot, and essay. AI auto-generates questions from policy content.",
           bold_prefix="Comprehension Quizzes \u2014")

add_bullet("Track document opens, time spent reading, device type, and access patterns",
           bold_prefix="Read Receipt Analytics \u2014")

add_bullet("Progressive escalation: friendly reminder \u2192 manager notification \u2192 "
           "auto-escalation with configurable SLA thresholds",
           bold_prefix="Automated Reminders \u2014")

add_bullet("Target by department, role, location, or security group with campaign-level "
           "analytics and completion tracking",
           bold_prefix="Distribution Campaigns \u2014")

add_bullet("30+ branded email templates delivered via Azure Logic App pipeline with retry "
           "and failure handling",
           bold_prefix="Email Notifications \u2014")

add_sub_heading("Analytics & Executive Reporting")

add_bullet("6-tab analytics centre covering executive overview, policy metrics, acknowledgements, "
           "SLA tracking, compliance risk, and audit reports",
           bold_prefix="Executive Dashboard \u2014")

add_bullet("Team compliance scores, pending approvals, review schedules, delegation management, "
           "and downloadable reports",
           bold_prefix="Manager Compliance View \u2014")

add_bullet("Policy lifecycle analytics, acknowledgement rates, review schedules, quiz statistics, "
           "and activity history",
           bold_prefix="Author Performance Reports \u2014")

add_bullet("Configurable targets for acknowledgement, approval, review, and authoring timelines "
           "with breach detection and alerts",
           bold_prefix="SLA Compliance Engine \u2014")

add_bullet("Every action recorded with user, timestamp, and context. Compliance-relevant actions "
           "flagged separately. Data export for external auditors.",
           bold_prefix="Complete Audit Log \u2014")


# ═══════════════════════════════════════════════════════════
# SECTION 5 — AI CAPABILITIES
# ═══════════════════════════════════════════════════════════

add_page_break()
add_section_header("05", "AI-Powered Intelligence",
    "The AI Advantage",
    "Powered by Azure OpenAI GPT-4o \u2014 secure, private, enterprise-grade")

add_body(
    "PolicyIQ includes a **conversational AI assistant** that understands your organisation\u2019s "
    "policies. Employees can ask natural-language questions and get instant, accurate answers with "
    "citations back to source policies \u2014 no more searching through documents or waiting for HR "
    "to respond."
)

add_sub_heading("Three Conversation Modes")

add_styled_table(
    ["Mode", "Purpose", "Example Questions"],
    [
        ["Policy Q&A",
         "Employees ask questions in plain English. The AI searches published policies, finds relevant content, and provides answers with direct citations.",
         "\u201CWhat is the return policy for pharmacy items?\u201D\n\u201CHow many sick leave days am I entitled to?\u201D"],
        ["Author Assistant",
         "Policy writers get AI-powered help drafting sections, improving clarity, checking for compliance gaps, and ensuring consistent language.",
         "\u201CDraft a disciplinary procedure section\u201D\n\u201CSimplify this legal language for employees\u201D"],
        ["Application Guide",
         "Contextual help for navigating PolicyIQ itself \u2014 how to submit for review, find reports, set up campaigns.",
         "\u201CHow do I approve a pending policy?\u201D\n\u201CWhere can I see team compliance?\u201D"],
    ],
    col_widths=[3.5, 6.5, 6]
)

add_sub_heading("AI Quiz Generation")

add_bullet("Reads full policy document content and generates comprehension questions automatically")
add_bullet("Configurable difficulty level, question count, and question types")
add_bullet("Authors review, edit, and approve before publishing to employees")
add_bullet("Supports all 11 question types (multiple choice, matching, ordering, essay, etc.)")

add_sub_heading("AI Document Classification")

add_bullet("Extracts text from uploaded DOCX, PDF, and PPTX documents")
add_bullet("Auto-classifies category and compliance risk level")
add_bullet("Suggests metadata including review frequency, tags, and key points")
add_bullet("Confidence scoring: Strong / Likely / Possible")
add_bullet("Batch processing for bulk imports of up to 50 documents")

add_quote(
    "All AI processing uses Azure OpenAI within your Microsoft Azure tenant. Policy content never "
    "leaves the Azure boundary. No data is used for model training. Enterprise-grade security and "
    "privacy by design."
)


# ═══════════════════════════════════════════════════════════
# SECTION 6 — USER EXPERIENCES
# ═══════════════════════════════════════════════════════════

add_page_break()
add_section_header("06", "User Experiences",
    "Tailored for Every Role",
    "17 purpose-built views across 4 access levels")

add_styled_table(
    ["Role", "Key Capabilities", "Primary Views"],
    [
        ["Employee",
         "Browse policies, acknowledge, take quizzes, search, get AI help, track personal compliance",
         "Start Screen, Policy Hub, My Policies, Policy Details, Search, Help"],
        ["Author",
         "Create & edit policies, manage drafts pipeline, handle approvals, build quizzes, run distribution",
         "Author Dashboard, Policy Builder, Quiz Builder, Policy Packs, Author Reports"],
        ["Manager",
         "Team compliance oversight, approve/delegate, request policies, review reports, manage distributions",
         "Manager Dashboard, Analytics, Distribution, Approvals, Delegations"],
        ["Admin",
         "System configuration, templates, workflows, SLA targets, user management, provisioning, diagnostics",
         "Admin Centre (21 sections), Event Viewer, Bulk Upload, all other views"],
    ],
    col_widths=[2.5, 6.5, 7]
)


# ═══════════════════════════════════════════════════════════
# SECTION 7 — BUSINESS SCENARIOS
# ═══════════════════════════════════════════════════════════

add_section_header("07", "Business Scenarios",
    "PolicyIQ in Action at Clicks",
    "Real-world scenarios demonstrating measurable value")

add_sub_heading("Scenario 1: New Store Opening \u2014 Onboarding 40 Employees")

add_body("Admin creates a **Policy Pack** containing all mandatory policies (Code of Conduct, "
         "H&S, Data Privacy, Pharmacy Protocols). The pack is assigned via a **Distribution Campaign** "
         "targeting the new store\u2019s security group. Employees receive branded email notifications, "
         "acknowledge each policy \u2014 some with **AI-generated comprehension quizzes** \u2014 and the "
         "Manager Dashboard shows real-time compliance progress. **100% within 5 days.** Complete "
         "audit trail stored for regulatory inspection.")

add_sub_heading("Scenario 2: Regulatory Update \u2014 Pharmacy Dispensing Protocol")

add_body("Author opens the existing policy and creates a new **revision** (original stays published). "
         "Updated content goes through **multi-level approval** (Pharmacy Manager \u2192 Compliance "
         "Officer \u2192 Legal). On publish, PolicyIQ auto-distributes to all pharmacy staff. "
         "**Automated reminders** escalate progressively. The Executive Dashboard shows national "
         "compliance climbing in real time. Non-compliant employees flagged for follow-up \u2014 "
         "**zero manual tracking**.")

add_sub_heading("Scenario 3: Annual Compliance Audit")

add_body("Auditor requests proof of policy acknowledgement for 12 months. Compliance Officer opens "
         "**Analytics Dashboard** and generates reports in seconds. Per-policy acknowledgement rates, "
         "individual employee histories, and complete audit trails exported. Quiz scores prove "
         "comprehension, not just checkbox compliance. Immutable, timestamped records \u2014 "
         "**audit passes with full evidence**.")


# ═══════════════════════════════════════════════════════════
# SECTION 8 — ROI
# ═══════════════════════════════════════════════════════════

add_page_break()
add_section_header("08", "Return on Investment",
    "Measurable Business Impact",
    "Efficiency gains, risk reduction, and cost savings")

add_kpi_table([
    ("60%", "Less Distribution Time"),
    ("80%", "Faster Ack Tracking"),
    ("50%", "Less Audit Prep"),
    ("100%", "Audit Trail Coverage"),
])

add_sub_heading("Efficiency Gains")
add_bullet("Eliminate manual distribution (email/paper)")
add_bullet("Zero spreadsheet tracking")
add_bullet("Automated reminder escalation")
add_bullet("AI-generated quizzes save author hours")
add_bullet("Bulk upload replaces one-by-one entry")
add_bullet("Self-service employee portal reduces HR queries")

add_sub_heading("Risk Mitigation")
add_bullet("Immediate non-compliance identification")
add_bullet("Provable acknowledgement evidence")
add_bullet("Comprehension testing (not just checkbox)")
add_bullet("SLA breach detection and alerts")
add_bullet("Regulatory-ready audit reports on demand")
add_bullet("Version control eliminates outdated policies")


# ═══════════════════════════════════════════════════════════
# SECTION 9 — COMPLIANCE & SECURITY
# ═══════════════════════════════════════════════════════════

add_section_header("09", "Compliance & Security",
    "Built for Regulated Industries",
    "Enterprise security, data protection, and regulatory readiness")

add_styled_table(
    ["Capability", "Description"],
    [
        ["POPIA Compliance", "PII redaction in telemetry, configurable data retention policies, right-to-access support"],
        ["Data Security", "Role-based access control, row-level security, encrypted storage via SharePoint Online"],
        ["Immutable Audit Trail", "Tamper-proof records, compliance-flagged events, complete action history with export"],
        ["SLA Monitoring", "Configurable targets for ack/approval/review timelines, breach detection, progressive alerts"],
        ["PII Protection", "Auto-redaction of emails and phone numbers in Application Insights telemetry"],
        ["Risk Classification", "5-level scoring: Critical, High, Medium, Low, Informational \u2014 per policy"],
    ],
    col_widths=[4, 12]
)


# ═══════════════════════════════════════════════════════════
# SECTION 10 — TECHNOLOGY
# ═══════════════════════════════════════════════════════════

add_page_break()
add_section_header("10", "Technology Platform",
    "Built on Microsoft 365",
    "No new platforms to learn, license, or manage")

add_body(
    "PolicyIQ is built entirely on the **Microsoft 365 ecosystem** using SharePoint Framework (SPFx), "
    "Azure Functions, and Azure OpenAI. This means Clicks can deploy PolicyIQ within its existing "
    "Microsoft tenant with **zero additional platform licensing**."
)

add_styled_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Frontend", "SharePoint Framework (SPFx) 1.20 + React + Fluent UI", "17 responsive webparts deployed to SharePoint Online"],
        ["Data Layer", "SharePoint Online Lists & Libraries", "30+ structured lists with full CRUD, versioning, and indexing"],
        ["AI Services", "Azure Functions + Azure OpenAI (GPT-4o)", "Policy chat, quiz generation, document classification"],
        ["Email Pipeline", "Azure Logic App + Office 365 Connector", "Reliable, branded email delivery with retry and status tracking"],
        ["Identity", "Microsoft Entra ID (Azure AD)", "SSO, role resolution, user sync, security groups"],
        ["Monitoring", "Application Insights + Event Viewer", "Real-time diagnostics, performance monitoring, error tracking"],
    ],
    col_widths=[3, 6.5, 6.5]
)


# ═══════════════════════════════════════════════════════════
# SECTION 11 — IMPLEMENTATION
# ═══════════════════════════════════════════════════════════

add_section_header("11", "Implementation Approach",
    "Deployment & Rollout",
    "A phased approach tailored to Clicks\u2019 environment")

add_styled_table(
    ["Phase", "Activities", "Duration"],
    [
        ["1. Discovery", "Stakeholder workshops, policy inventory, role mapping, compliance requirements analysis", "1\u20132 weeks"],
        ["2. Configuration", "Tenant setup, SharePoint provisioning, Azure resource deployment, branding customisation, Entra ID sync", "1\u20132 weeks"],
        ["3. Migration", "Bulk import existing policies, metadata assignment, AI classification, template creation", "2\u20133 weeks"],
        ["4. UAT & Training", "User acceptance testing, role-based training workshops, admin handover, help centre content", "1\u20132 weeks"],
        ["5. Go-Live", "Phased rollout (head office \u2192 regional \u2192 stores), hypercare support, feedback loop", "2\u20134 weeks"],
    ],
    col_widths=[3, 9.5, 3.5]
)

add_quote(
    "PolicyIQ is production-ready today. The platform has been through 24 development sessions, "
    "33 production hardening fixes, 164 end-to-end tests, and a comprehensive security audit. "
    "Deployment to Clicks is configuration \u2014 not development."
)


# ═══════════════════════════════════════════════════════════
# SECTION 12 — PRICING
# ═══════════════════════════════════════════════════════════

add_page_break()
add_section_header("12", "Investment",
    "Pricing",
    "A turnkey solution \u2014 fully configured, deployed, and supported")

# Price callout
add_blank(1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("R175,000")
run.font.name = 'Segoe UI'
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Once-off implementation fee  \u2014  excludes VAT")
run.font.name = 'Segoe UI'
run.font.size = Pt(11)
run.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(20)

add_sub_heading("What\u2019s Included")

add_bullet("Full PolicyIQ platform deployment to Clicks\u2019 Microsoft 365 tenant")
add_bullet("SharePoint Online provisioning (30+ lists, libraries, and pages)")
add_bullet("Azure resource deployment (AI services, email pipeline, monitoring)")
add_bullet("Branding customisation to Clicks\u2019 corporate identity")
add_bullet("Entra ID integration and role mapping")
add_bullet("Bulk migration of existing policies (up to 200 documents)")
add_bullet("AI-powered quiz generation and document classification")
add_bullet("Conversational AI policy assistant (3 modes)")
add_bullet("Role-based training workshops (Admin, Author, Manager, Employee)")
add_bullet("30 days post-go-live hypercare support")
add_bullet("Complete technical documentation and admin handover")

add_sub_heading("Terms")

add_styled_table(
    ["Item", "Detail"],
    [
        ["Payment Terms", "50% on acceptance of this proposal. 50% on successful go-live sign-off. Invoiced in South African Rand (ZAR)."],
        ["Ongoing Support", "Optional annual support and maintenance agreement available at 15% of implementation fee. Includes updates, bug fixes, and priority support."],
        ["Azure Costs", "Azure consumption costs (AI services, Logic App, Functions) are billed directly to Clicks\u2019 Azure subscription. Estimated R1,500 \u2013 R3,000/month."],
        ["Validity", "This proposal is valid for 30 days from the date of issue. Pricing subject to review after this period."],
        ["M365 Licensing", "No additional Microsoft licensing required. PolicyIQ runs entirely on Clicks\u2019 existing Microsoft 365 E3/E5 subscription and SharePoint Online entitlements."],
    ],
    col_widths=[3.5, 12.5]
)


# ═══════════════════════════════════════════════════════════
# SECTION 13 — DWx ECOSYSTEM
# ═══════════════════════════════════════════════════════════

add_section_header("13", "The DWx Ecosystem",
    "Part of Something Bigger",
    "PolicyIQ is one application in the DWx Digital Workplace Excellence suite")

add_body(
    "PolicyIQ operates as a **standalone application**, but it is designed to integrate with the "
    "broader DWx suite. As Clicks\u2019 digital workplace matures, additional DWx applications can "
    "be deployed alongside PolicyIQ \u2014 each sharing common components, navigation patterns, and "
    "an optional cross-app hub for unified notifications and activity feeds."
)

add_styled_table(
    ["Application", "Description", "Status"],
    [
        ["PolicyIQ", "Policy governance, compliance tracking, acknowledgement management, AI-powered intelligence", "Production Ready"],
        ["Asset Manager", "IT asset lifecycle, vendor management, procurement, maintenance tracking", "Available"],
        ["Contract Manager", "Contract lifecycle, renewals, obligations tracking, vendor compliance", "Available"],
        ["Future Apps", "Additional DWx applications planned for the suite based on client needs", "Roadmap"],
    ],
    col_widths=[3.5, 9, 3.5]
)


# ═══════════════════════════════════════════════════════════
# SECTION 14 — SIGNING PAGE
# ═══════════════════════════════════════════════════════════

add_page_break()
add_section_header("14", "Acceptance",
    "Proposal Acceptance",
    "Agreement to proceed with PolicyIQ implementation")

add_body(
    "By signing below, both parties agree to the terms outlined in this proposal for the "
    "implementation of **PolicyIQ** for Clicks. This agreement covers the scope of work detailed "
    "in Sections 1\u201313, including all deliverables, timelines, and the investment of "
    "**R175,000 (excl. VAT)** as specified in Section 12."
)

add_blank(1)

# ── Client signing ──
p = doc.add_paragraph()
run = p.add_run("ON BEHALF OF CLICKS")
run.font.name = 'Segoe UI'
run.font.size = Pt(10)
run.font.color.rgb = TEAL
run.bold = True
p.paragraph_format.space_after = Pt(4)
# Top border
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="8" w:space="4" w:color="0F172A"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

add_signing_field("Full Name")
add_signing_field("Title / Position")
add_signing_field("Signature")
add_signing_field("Date")

add_blank(2)

# ── Provider signing ──
p = doc.add_paragraph()
run = p.add_run("ON BEHALF OF FIRST DIGITAL")
run.font.name = 'Segoe UI'
run.font.size = Pt(10)
run.font.color.rgb = TEAL
run.bold = True
p.paragraph_format.space_after = Pt(4)
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="8" w:space="4" w:color="0F172A"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

add_signing_field("Full Name")
add_signing_field("Title / Position")
add_signing_field("Signature")
add_signing_field("Date")

add_blank(2)

# Terms note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run("Note: ")
run.font.name = 'Segoe UI'
run.font.size = Pt(9)
run.font.color.rgb = TEAL_DARK
run.bold = True
run2 = p.add_run(
    "This proposal is valid for 30 days from the date of issue (April 2026). "
    "Payment terms: 50% on acceptance, 50% on go-live sign-off. "
    "All amounts are in South African Rand (ZAR) and exclude VAT at the prevailing rate."
)
run2.font.name = 'Segoe UI'
run2.font.size = Pt(9)
run2.font.color.rgb = MUTED


# ═══════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════

add_blank(2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PolicyIQ")
run.font.name = 'Segoe UI'
run.font.size = Pt(9)
run.font.color.rgb = TEAL
run.bold = True
run2 = p.add_run(" by First Digital \u2014 DWx Digital Workplace Excellence Suite")
run2.font.name = 'Segoe UI'
run2.font.size = Pt(9)
run2.font.color.rgb = MUTED

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p2.add_run("This document is confidential and intended solely for Clicks. \u00A9 2026 First Digital. All rights reserved.")
run3.font.name = 'Segoe UI'
run3.font.size = Pt(8)
run3.font.color.rgb = MUTED


# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), "PolicyIQ-Proposal-Clicks.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
print(f"Size: {os.path.getsize(output_path) / 1024:.0f} KB")
